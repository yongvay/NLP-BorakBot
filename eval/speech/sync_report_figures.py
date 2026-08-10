"""Rewrite the report's result tables from the result CSVs.

WHY THIS EXISTS
---------------
Table A in PartB_Draft_STT.docx was transcribed by hand from console output across
several evaluation runs. Three of its nine rows then drifted away from the CSVs backing
them -- the vanilla configurations were re-transcribed after the numbers were copied,
and nobody re-copied them. The report and its own evidence disagreed by up to 0.028.

Hand-copying numbers into a document is the failure; doing it more carefully is not the
fix. This script makes the CSVs the single source of truth and the document a rendering
of them. Run it after every evaluation run.

WHAT IT DOES
------------
    tables  rewritten in place from results_summary.csv / results_by_category.csv
    prose   NOT rewritten -- figures embedded in sentences are checked and reported,
            because a number in prose usually carries a claim ("22 points", "a further
            6") that only a human can restate correctly.

Usage
    python eval/speech/sync_report_figures.py            # check only, no writes
    python eval/speech/sync_report_figures.py --write    # rewrite the tables
"""

import argparse
import re
import shutil
import sys
from pathlib import Path

import pandas as pd

HERE = Path(__file__).resolve().parent
REPO = HERE.parent.parent
DOCX = REPO / "docs" / "PartB_Draft_STT.docx"

# Table A labels as typed in the report -> config key in the CSVs.
TABLE_A_LABELS = {
    "malaysian_prompt_routed": "malaysian_prompt_routed",
    "malaysian_prompt_auto": "malaysian_prompt_auto",
    "malaysian_prompt (forced ms)": "malaysian_prompt",
    "vanilla_prompt_routed": "vanilla_prompt_routed",
    "vanilla_prompt": "vanilla_prompt",
    "vanilla_noprompt": "vanilla_noprompt",
    "malaysian_noprompt_auto": "malaysian_noprompt_auto",
    "malaysian_noprompt (forced ms)": "malaysian_noprompt",
    "malaysian_prompt_en (diagnostic)": "malaysian_prompt_en",
}

BEST = "malaysian_prompt_routed"


def _find(name: str) -> Path:
    for p in (HERE / name, HERE / "results" / name):
        if p.exists():
            return p
    sys.exit(f"missing {name}")


def _set_cell(cell, value: str) -> bool:
    """Replace a cell's text, keeping the first run's formatting. True if changed."""
    para = cell.paragraphs[0]
    if not para.runs:
        para.add_run(value)
        return True
    if para.runs[0].text == value and len(para.runs) == 1:
        return False
    para.runs[0].text = value
    for r in para.runs[1:]:
        r.text = ""
    return True


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--write", action="store_true", help="apply changes (default: check only)")
    args = ap.parse_args()

    import docx

    summary = pd.read_csv(_find("results_summary.csv")).set_index("config")
    cat = pd.read_csv(_find("results_by_category.csv"))
    doc = docx.Document(DOCX)

    changes = []

    # ---------------------------------------------------------------- Table A
    for t in doc.tables:
        head = [c.text.strip() for c in t.rows[0].cells]
        if head[:2] != ["Configuration", "WER (strict)"]:
            continue
        for row in t.rows[1:]:
            label = row.cells[0].text.strip()
            key = TABLE_A_LABELS.get(label)
            if key is None:
                print(f"  ! unrecognised Table A row: {label!r}")
                continue
            if key not in summary.index:
                print(f"  ! {key} missing from results_summary.csv")
                continue
            s = summary.loc[key]
            want = [f"{s.wer_strict:.3f}", f"{s.wer_lenient:.3f}",
                    str(int(s['sub'])), str(int(s['ins'])), str(int(s['dele']))]
            for ci, v in enumerate(want, start=1):
                old = row.cells[ci].text.strip()
                if old != v:
                    changes.append(f"Table A  {label:34} col{ci}  {old!r} -> {v!r}")
                    if args.write:
                        _set_cell(row.cells[ci], v)

    # ---------------------------------------------------------------- Table B
    b = cat[cat.config == BEST].sort_values("wer_strict")
    by_cat = {r.switch_type: r for r in b.itertuples()}
    overall = summary.loc[BEST]
    for t in doc.tables:
        head = [c.text.strip() for c in t.rows[0].cells]
        if head[:2] != ["Category", "Reference words"]:
            continue
        for row in t.rows[1:]:
            label = row.cells[0].text.strip()
            if label.lower() == "overall":
                want = [str(int(overall.ref_words)), f"{overall.wer_strict:.3f}"]
            elif label in by_cat:
                r = by_cat[label]
                want = [str(int(r.ref_words)), f"{r.wer_strict:.3f}"]
            else:
                print(f"  ! unrecognised Table B row: {label!r}")
                continue
            for ci, v in enumerate(want, start=1):
                old = row.cells[ci].text.strip()
                if old != v:
                    changes.append(f"Table B  {label:34} col{ci}  {old!r} -> {v!r}")
                    if args.write:
                        _set_cell(row.cells[ci], v)

    # ------------------------------------------------- prose figures (report only)
    # Every 0.xxx in the body is checked against the set of values the CSVs actually
    # contain. A figure that matches nothing is either stale or a number this script
    # does not know about -- both are worth a human look, neither is safe to auto-edit.
    known = set()
    for v in summary.wer_strict:
        known.add(f"{v:.3f}")
    for v in summary.wer_lenient:
        known.add(f"{v:.3f}")
    for v in cat.wer_strict:
        known.add(f"{v:.3f}")

    # detector probabilities and the threshold sweep, if those runs have been done
    for name, cols in (("detector_probs.csv", ["p_en"]),
                       ("threshold_sweep.csv", ["threshold", "wer_strict", "wer_lenient"])):
        for p in (HERE / name, HERE / "results" / name):
            if p.exists():
                extra = pd.read_csv(p)
                for c in cols:
                    for v in extra[c]:
                        known.add(f"{v:.3f}")
                        known.add(f"{v:.4f}")
                # per-category min/median/max are quoted in the sensitivity section
                if "switch_type" in extra.columns and "p_en" in extra.columns:
                    g = extra.groupby("switch_type").p_en
                    for series in (g.min(), g.median(), g.max()):
                        for v in series:
                            known.add(f"{v:.3f}")
                break

    known |= {"0.250", "0.300", "0.500", "0.277"}   # target, cited literature range, oracle

    suspicious = []
    for i, p in enumerate(doc.paragraphs):
        for m in re.findall(r"\b0\.\d{3}\b", p.text):
            if m not in known:
                suspicious.append((i, m, p.text[:110]))

    # ---------------------------------------------------------------- report
    print(f"document: {DOCX.name}")
    print(f"tables:   {len(changes)} cell(s) out of sync with the CSVs")
    for c in changes:
        print("   ", c)
    if suspicious:
        print(f"\nprose:    {len(suspicious)} figure(s) matching no value in the CSVs "
              f"-- check by hand:")
        for i, m, ctx in suspicious:
            print(f"    para {i}: {m}   {ctx}...")
    else:
        print("prose:    every 0.xxx figure matches a value in the CSVs")

    if not args.write:
        print("\n(check only -- pass --write to apply table changes)")
        return 1 if changes else 0

    if changes:
        backup = DOCX.with_suffix(".docx.bak")
        shutil.copy2(DOCX, backup)
        doc.save(DOCX)
        print(f"\nwrote {DOCX.name}  (backup: {backup.name})")
    else:
        print("\nnothing to do")
    return 0


if __name__ == "__main__":
    sys.exit(main())
